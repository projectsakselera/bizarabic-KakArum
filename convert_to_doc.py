#!/usr/bin/env python3
"""
Convert Markdown to Word Document (.docx)
Preserves formatting, Arabic text, and structure
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_md_to_docx(md_file, output_file):
    """Convert markdown file to Word document"""

    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create a new Document
    doc = Document()

    # Split content by lines
    lines = content.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)

        # Handle HTML divs with Arabic text (RTL)
        elif line.startswith('<div dir="rtl"'):
            # Collect all lines until closing div
            div_content = []
            i += 1
            while i < len(lines) and not lines[i].startswith('</div>'):
                div_line = lines[i]
                # Remove HTML tags
                div_line = re.sub(r'<[^>]+>', '', div_line)
                if div_line.strip():
                    div_content.append(div_line)
                i += 1

            # Add as right-aligned paragraph
            if div_content:
                p = doc.add_paragraph('\n'.join(div_content))
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.size = Pt(12)

        # Handle tables
        elif line.startswith('| ') or line.startswith('|--'):
            table_lines = []
            while i < len(lines) and (lines[i].startswith('| ') or lines[i].startswith('|--')):
                if not lines[i].startswith('|--'):  # Skip separator lines
                    table_lines.append(lines[i])
                i += 1
            i -= 1  # Step back one

            if table_lines:
                # Parse table
                rows = []
                for tline in table_lines:
                    cells = [cell.strip() for cell in tline.split('|')[1:-1]]
                    rows.append(cells)

                if rows:
                    # Create table
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'

                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            # Remove markdown formatting
                            cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_data)
                            cell_text = re.sub(r'<[^>]+>', '', cell_text)
                            table.rows[row_idx].cells[col_idx].text = cell_text

        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Remove markdown formatting
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')

        # Handle numbered lists
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')

        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph('_' * 50)

        # Handle code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1

            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines))
                p.style = 'Intense Quote'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)

        # Regular paragraph
        else:
            # Remove HTML tags and markdown formatting
            text = re.sub(r'<[^>]+>', '', line)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)  # Remove links
            text = re.sub(r'`(.+?)`', r'\1', text)

            if text.strip():
                # Check if contains Arabic (right-to-left)
                if re.search(r'[\u0600-\u06FF]', text):
                    p = doc.add_paragraph(text)
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p = doc.add_paragraph(text)

        i += 1

    # Save the document
    doc.save(output_file)
    print(f"✅ Successfully converted to: {output_file}")

if __name__ == "__main__":
    input_file = "01-MODULE-SESSION-1.md"
    output_file = "01-MODULE-SESSION-1.docx"

    convert_md_to_docx(input_file, output_file)
